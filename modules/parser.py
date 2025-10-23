import pdfplumber
from docx import Document
import io

def extract_text_from_pdf(file):
    """
    file: can be a path (str) or a Streamlit UploadedFile object
    """
    text = ""
    if hasattr(file, "read"):  # UploadedFile
        with pdfplumber.open(io.BytesIO(file.read())) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    else:  # file path
        with pdfplumber.open(file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    return text

def extract_text_from_docx(file):
    """
    file: can be a path (str) or a Streamlit UploadedFile object
    """
    if hasattr(file, "read"):  # UploadedFile
        file_stream = io.BytesIO(file.read())
        doc = Document(file_stream)
    else:  # file path
        doc = Document(file)
    return "\n".join([para.text for para in doc.paragraphs])

def extract_text(file):
    """
    Automatically detect file type and extract text.
    Supports PDF and DOCX.
    """
    filename = getattr(file, "name", "")  # UploadedFile has .name
    if isinstance(file, str):
        filename = file

    if filename.lower().endswith(".pdf"):
        return extract_text_from_pdf(file)
    elif filename.lower().endswith(".docx"):
        return extract_text_from_docx(file)
    else:
        return ""
